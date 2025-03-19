import PyPDF2
import docx
import cv2
import pytesseract
import re
import string
import numpy as np
import streamlit as st
import pandas as pd
from datetime import datetime
from sentence_transformers import SentenceTransformer, util
import plotly.graph_objects as go
import plotly.express as px
import time
import concurrent.futures
from collections import Counter


@st.cache_resource(show_spinner=False)
def load_model():
    return SentenceTransformer('all-MiniLM-L6-v2')

@st.cache_data(show_spinner=False)
def process_resume(file_content, file_type):
    if file_type == "application/pdf":
        return extract_text_from_pdf(file_content)
    elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
        return extract_text_from_docx(file_content)
    elif file_type.startswith("image"):
        return extract_text_from_image(file_content)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

st_model = load_model()


def extract_text_from_pdf(file):
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        if len(pdf_reader.pages) == 0:
            raise ValueError("PDF file appears to be empty")
        
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + " "
        
        if not text.strip():
            raise ValueError("No extractable text found in PDF")
        return text
    except PyPDF2.errors.PdfReadError as e:
        st.error(f"Invalid or corrupted PDF file: {e}")
        return ""
    except Exception as e:
        st.error(f"Error extracting PDF: {e}")
        return ""

def extract_text_from_docx(file):
    try:
        doc = docx.Document(file)
        if len(doc.paragraphs) == 0:
            raise ValueError("DOCX file appears to be empty")
        
        full_text = [para.text for para in doc.paragraphs]
        text = " ".join(full_text)
        
        if not text.strip():
            raise ValueError("No extractable text found in DOCX")
        return text
    except Exception as e:
        st.error(f"Error extracting DOCX: {e}")
        return ""

def extract_text_from_image(file):
    try:
        file_bytes = np.asarray(bytearray(file.read()), dtype=np.uint8)
        img = cv2.imdecode(file_bytes, cv2.IMREAD_COLOR)
        
        if img is None or img.size == 0:
            raise ValueError("Invalid image file or empty image")
        
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        _, binary = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)
        denoised = cv2.fastNlMeansDenoising(binary, None, 10, 7, 21)
        
        text = pytesseract.image_to_string(denoised)
        
        if not text.strip():
            raise ValueError("No text could be extracted from the image")
        return text
    except Exception as e:
        st.error(f"Error extracting image: {e}")
        return ""

def basic_preprocess(text):
    text = re.sub(r'\s+', ' ', text).strip()
    text = re.sub(r'\n+', '\n', text)
    processed_text = text.lower()
    translator = str.maketrans('', '', string.punctuation)
    processed_text = processed_text.translate(translator)
    return processed_text


def extract_keywords(text, job_description):
    words = re.findall(r'\b[a-zA-Z][a-zA-Z0-9+#]{2,}\b', text.lower())
    job_words = re.findall(r'\b[a-zA-Z][a-zA-Z0-9+#]{2,}\b', job_description.lower())
    
    word_counts = Counter(words)
    job_word_counts = Counter(job_words)
    
    top_resume_words = dict(word_counts.most_common(20))
    
    matching_keywords = set(top_resume_words.keys()) & set(job_word_counts.keys())
    
    return list(matching_keywords), list(top_resume_words.keys())


def compute_similarity(resume_text, job_description):
    resume_embedding = st_model.encode(resume_text, convert_to_tensor=True)
    job_embedding = st_model.encode(job_description, convert_to_tensor=True)
    similarity_score = util.cos_sim(resume_embedding, job_embedding)
    return similarity_score.item()

def compute_final_score(base_score, boost_points, keyword_match_score=0, boost_weight=0.05):
    keyword_component = keyword_match_score * 0.1
    final_score = base_score + keyword_component + (boost_points * boost_weight)
    return min(final_score, 1.0)

def compute_score_breakdown(base_score, boost_points, keyword_match_score):
    breakdown = {
        "Skills Matching": round(0.35 * base_score, 4),
        "Experience Relevance": round(0.35 * base_score, 4),
        "Education Relevance": round(0.15 * base_score, 4),
        "Keyword Matching": round(keyword_match_score * 0.1, 4),
        "Boost Points": round(boost_points * 0.05, 4)
    }
    return breakdown

def generate_improvement_suggestions(final_score, breakdown, text_length, matching_keywords, job_description):
    suggestions = []
    
    if final_score < 0.5:
        suggestions.append("Your resume needs significant improvement. Consider adding more relevant skills and experiences.")
    elif final_score < 0.7:
        suggestions.append("Your resume shows potential but could be improved. Tailor it more closely to match the job requirements.")
    elif final_score < 0.9:
        suggestions.append("Your resume is strong! Consider highlighting your key achievements more prominently.")
    else:
        suggestions.append("Excellent resume! Just minor tweaks could make it even better.")
    
    if text_length < 200:
        suggestions.append("Your resume seems too brief. Consider adding more details about your responsibilities and achievements.")
    elif text_length > 800:
        suggestions.append("Your resume is quite lengthy. Consider focusing on the most relevant experiences and skills.")
    
    if len(matching_keywords) < 5:
        suggestions.append("Your resume lacks key terms from the job description. Consider adding more relevant keywords.")
    
    job_skills = re.findall(r'\b(?:experience|proficiency|knowledge|familiarity)\s+(?:with|in|of)?\s+([A-Za-z0-9+#]+)', job_description)
    if job_skills:
        missing_skills = [skill for skill in job_skills if skill.lower() not in ' '.join(matching_keywords).lower()]
        if missing_skills and len(missing_skills) <= 5:
            suggestions.append(f"Consider highlighting your experience with: {', '.join(missing_skills[:5])}")
    
    return suggestions


def generate_report(candidate_info, preprocessed_text, job_description):
    report = f"Resume Analysis Report - {candidate_info['Candidate']}\n"
    report += f"Processed on: {candidate_info['Timestamp']}\n{'-'*50}\n\n"
    
    report += "SUMMARY\n"
    report += f"Base Similarity Score: {candidate_info['Base Score']:.4f}\n"
    report += f"Keyword Match Score: {candidate_info['Keyword Match Score']:.4f}\n"
    report += f"Manual Boost Points: {candidate_info['Boost Points']}\n"
    report += f"Final Score: {candidate_info['Final Score']:.4f}\n"
    report += f"Word Count: {candidate_info['Word Count']}\n\n"
    
    report += "SCORE BREAKDOWN\n"
    for key, value in candidate_info['Breakdown'].items():
        report += f"  {key}: {value:.4f}\n"
    
    report += "\nKEYWORDS ANALYSIS\n"
    report += f"Matching Keywords: {', '.join(candidate_info['Matching Keywords'])}\n"
    report += f"Top Resume Terms: {', '.join(candidate_info['Top Keywords'][:10])}\n\n"
    
    report += "IMPROVEMENT SUGGESTIONS\n"
    for i, s in enumerate(candidate_info['Suggestions'], 1):
        report += f"{i}. {s}\n"
    
    report += "\nREFERENCE\n"
    report += "Job Description:\n" + job_description + "\n\n"
    report += "Preprocessed Resume Text:\n" + preprocessed_text + "\n"
    
    return report


def display_candidate_results(candidate_info, i):
    st.markdown(f"### {candidate_info['Candidate']}")
    
    col1, col2 = st.columns(2)
    with col1:
        fig = go.Figure(go.Indicator(
            mode="gauge+number",
            value=candidate_info['Final Score'] * 100,
            domain={'x': [0, 1], 'y': [0, 1]},
            title={'text': "Final Score (%)"},
            gauge={
                'axis': {'range': [0, 100]},
                'bar': {'color': "#1f77b4"},
                'steps': [
                    {'range': [0, 50], 'color': "lightgray"},
                    {'range': [50, 75], 'color': "gray"},
                    {'range': [75, 100], 'color': "lightgreen"}
                ],
                'threshold': {
                    'line': {'color': "red", 'width': 4},
                    'thickness': 0.75,
                    'value': candidate_info['Final Score'] * 100
                }
            }
        ))
        st.plotly_chart(fig, use_container_width=True)
    
    with col2:
        st.subheader("Improvement Suggestions")
        for idx, suggestion in enumerate(candidate_info['Suggestions']):
            st.write(f"- {suggestion}", key=f"cand_{i}_suggestion_{idx}")
    
    with st.expander("Detailed Analysis"):
        st.subheader("Score Breakdown")
        breakdown_df = pd.DataFrame({
            'Component': list(candidate_info['Breakdown'].keys()),
            'Score': list(candidate_info['Breakdown'].values())
        })
        fig_breakdown = px.bar(breakdown_df, x='Component', y='Score', 
                              title="Score Components", color='Score')
        st.plotly_chart(fig_breakdown, use_container_width=True)
        
        st.subheader("Keyword Analysis")
        st.write("**Matching Keywords:**", ", ".join(candidate_info['Matching Keywords']))
        st.write("**Top Resume Terms:**", ", ".join(candidate_info['Top Keywords'][:10]))
        
        st.write("**Candidate Name:**", candidate_info['Candidate'])
        st.write("**Base Score:**", round(candidate_info['Base Score'], 4))
        st.write("**Keyword Match Score:**", round(candidate_info['Keyword Match Score'], 4))
        st.write("**Boost Points:**", candidate_info['Boost Points'])
        st.write("**Word Count:**", candidate_info['Word Count'])
        st.write("**Timestamp:**", candidate_info['Timestamp'])
        
        report_text = generate_report(candidate_info, candidate_info['Preprocessed Text'], candidate_info['Job Description'])
        st.download_button("Download Candidate Report", report_text, 
                          file_name=f"{candidate_info['Candidate']}_Report.txt", 
                          mime="text/plain", key=f"download_{i}")

def display_comparison(candidate_results):
    if not candidate_results:
        return
        
    df = pd.DataFrame(candidate_results)
    df_sorted = df.sort_values(by="Final Score", ascending=False).reset_index(drop=True)
    best_candidate = df_sorted.iloc[0]["Candidate"]
    
    st.subheader("Candidate Ranking Comparison")
    st.write("The best candidate is **" + best_candidate + "** based on the final score.")
    
    comparison_df = df_sorted[["Candidate", "Base Score", "Keyword Match Score", 
                               "Boost Points", "Final Score", "Word Count"]]
    st.dataframe(comparison_df)
    
    fig_bar = px.bar(df_sorted, x="Candidate", y="Final Score", color="Final Score",
                    title="Final Score Comparison", text_auto=True)
    fig_bar.update_traces(texttemplate='%{y:.4f}', textposition='outside')
    st.plotly_chart(fig_bar, use_container_width=True)

    
    st.download_button(
        comparison_df.to_csv(index=False).encode('utf-8'),
        file_name="resume_comparison_results.csv",
        mime="text/csv"
    )


def process_candidate(file, job_description, i):
    try:
        progress_text = f"Processing {file.name}..."
        progress_bar = st.progress(0, text=progress_text)
        
        progress_bar.progress(10, text=f"Extracting text from {file.name}...")
        file.seek(0)
        file_type = file.type
        candidate_text = process_resume(file, file_type)
        
        if not candidate_text:
            st.error(f"Failed to extract text from {file.name}")
            progress_bar.empty()
            return None
            
        progress_bar.progress(30, text="Preprocessing text...")
        candidate_name = file.name.split('.')[0]
        preprocessed_text = basic_preprocess(candidate_text)
        
        progress_bar.progress(50, text="Computing similarity...")
        base_score = compute_similarity(preprocessed_text, job_description)
        
        progress_bar.progress(70, text="Extracting keywords...")
        matching_keywords, top_keywords = extract_keywords(preprocessed_text, job_description)
        keyword_match_score = len(matching_keywords) / max(len(top_keywords), 1) if top_keywords else 0
        
        progress_bar.progress(90, text="Generating suggestions...")
        boost_points = st.slider(f"Boost Points for {candidate_name}", 
                                min_value=0, max_value=10, value=0, step=1, 
                                key=f"boost_{i}")
        
        final_score = compute_final_score(base_score, boost_points, keyword_match_score)
        breakdown = compute_score_breakdown(base_score, boost_points, keyword_match_score)
        text_length = len(preprocessed_text.split())
        suggestions = generate_improvement_suggestions(final_score, breakdown, 
                                                      text_length, matching_keywords, 
                                                      job_description)
        
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        candidate_info = {
            "Candidate": candidate_name,
            "Base Score": round(base_score, 4),
            "Keyword Match Score": round(keyword_match_score, 4),
            "Boost Points": boost_points,
            "Final Score": round(final_score, 4),
            "Word Count": text_length,
            "Breakdown": breakdown,
            "Suggestions": suggestions,
            "Matching Keywords": matching_keywords,
            "Top Keywords": top_keywords,
            "Timestamp": timestamp,
            "Preprocessed Text": preprocessed_text,
            "Job Description": job_description
        }
        
        progress_bar.progress(100, text="Complete!")
        time.sleep(0.5)
        progress_bar.empty()
        
        return candidate_info
        
    except Exception as e:
        st.error(f"Error processing {file.name}: {e}")
        return None


def main():
    st.set_page_config(
        page_title="ResumeRise - Empower Your Career", 
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    st.title("ResumeRise")
    st.header("Resume Analysis & Comparison Tool")
    
    st.markdown("""
    ResumeRise helps you analyze and compare candidate resumes against job descriptions. 
    Upload multiple resumes, paste a job description, and get detailed insights on each candidate.
    
    **Features:**
    - Extract text from PDF, DOCX, and image files
    - Compute similarity between resumes and job descriptions
    - Extract and match keywords
    - Generate personalized improvement suggestions
    - Compare candidates side-by-side
    - Export comparison results as CSV
    """)
    
    with st.sidebar:
        st.header("Settings")
        st.write("Upload multiple candidate resumes and paste the job description.")
        
        uploaded_files = st.file_uploader(
            "Upload Candidate Resumes", 
            type=["pdf", "docx", "jpg", "png"], 
            accept_multiple_files=True, 
            key="file_uploader"
        )
        
        job_description = st.text_area("Paste the Job Description", height=200, key="job_desc")
        
        st.subheader("Analysis Settings")
        boost_enabled = st.checkbox("Enable Manual Boost Points", value=True)
        
        process_button = st.button("Process Resumes", type="primary", key="process_button")
    
    if uploaded_files and job_description and process_button:
        st.subheader("Processing Resumes...")
        
        if not all(file.size > 0 for file in uploaded_files):
            st.error("One or more files appear to be empty. Please check your uploads.")
            return
        
        total_files = len(uploaded_files)
        overall_progress = st.progress(0, text=f"Processing 0/{total_files} files...")
        
        candidate_results = []
        
        for i, file in enumerate(uploaded_files):
            candidate_info = process_candidate(file, job_description, i)
            if candidate_info:
                candidate_results.append(candidate_info)
                display_candidate_results(candidate_info, i)
                st.markdown("---")
            
            overall_progress.progress((i + 1) / total_files, text=f"Processing {i + 1}/{total_files} files...")
        
        if len(candidate_results) > 1:
            display_comparison(candidate_results)
        
        if not candidate_results:
            st.warning("No valid resumes were processed. Please check your files and try again.")
        
        overall_progress.empty()
    
    elif not uploaded_files:
        st.info("👈 Upload resumes and paste a job description to get started!")
        
        with st.expander("See an example result"):
            st.write("""
            Example output shows:
            - Candidate score gauge
            - Personalized improvement suggestions
            - Detailed breakdown of scores
            - Keyword matching analysis
            - Comparison between candidates
            """)

if __name__ == "__main__":
    main()